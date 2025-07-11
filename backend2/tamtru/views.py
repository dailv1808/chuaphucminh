# views.py
from docx import Document
import os
import threading
import mimetypes
from django.http import FileResponse
from rest_framework.decorators import api_view
from rest_framework.response import Response
from django.conf import settings

@api_view(['POST'])
def generate_tam_tru(request):
    try:
        # Lấy dữ liệu từ frontend
        data = request.data

        # Đường dẫn tới template file
        template_path = os.path.join(settings.BASE_DIR, 'tamtru', 'file', 'CT01.docx')
        if not os.path.exists(template_path):
            return Response({'error': 'Template file not found'}, status=404)

        # Mở template
        doc = Document(template_path)

        # Thay thế dữ liệu
        replacements = {
            '{{fullname}}': data.get('fullname', ''),
            '{{birthday}}': data.get('birthday', ''),
            '{{gender}}': data.get('gender', ''),
            '{{cccd}}': data.get('cccd', ''),
            '{{username}}': data.get('username', ''),
            '{{email}}': data.get('email', ''),
            '{{address}}': data.get('address', '')
        }

        # Thay thế trong các đoạn văn
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        # Thay thế trong bảng
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

        # Tạo thư mục nếu chưa tồn tại
        output_dir = os.path.join(settings.MEDIA_ROOT, 'tam_tru_docs')
        os.makedirs(output_dir, exist_ok=True)

        # Lưu file Word
        output_filename = f"CT01_{data.get('cccd', '')}.docx"
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)

        # Xác định content-type
        content_type, _ = mimetypes.guess_type(output_path)
        content_type = content_type or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        # Tạo response để trả file
        response = FileResponse(
            open(output_path, 'rb'),
            as_attachment=True,
            filename=output_filename,
            content_type=content_type
        )

        # Xóa file sau khi gửi response bằng thread
        def delete_file_later(path):
            import time
            time.sleep(5)  # Đợi vài giây để đảm bảo file đã được gửi xong
            if os.path.exists(path):
                try:
                    os.remove(path)
                    print(f"Đã xóa file tạm: {path}")
                except Exception as e:
                    print(f"Lỗi khi xóa file: {e}")

        threading.Thread(target=delete_file_later, args=(output_path,)).start()

        return response

    except Exception as e:
        return Response({'error': str(e)}, status=500)

